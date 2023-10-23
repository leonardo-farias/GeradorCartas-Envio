import openpyxl
from docx import Document
import datetime
from docx2pdf import convert
import locale  # Importe o módulo locale

# Defina a formatação de números para o padrão brasileiro
locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

# Variável data
data_atual = datetime.date.today()

# Formata a data no formato brasileiro (dd/mm/aaaa)
data_formatada = data_atual.strftime('%d/%m/%Y')

meses_extenso = {
    1: 'janeiro',
    2: 'fevereiro',
    3: 'março',
    4: 'abril',
    5: 'maio',
    6: 'junho',
    7: 'julho',
    8: 'agosto',
    9: 'setembro',
    10: 'outubro',
    11: 'novembro',
    12: 'dezembro'
}

data_extenso = f"{data_atual.day} de {meses_extenso[data_atual.month]} de {data_atual.year}"

# Carregar o arquivo Excel
workbook = openpyxl.load_workbook('C:/Users/leonardo.farias/Desktop/Projetos/Gerador de Cartas/Modelo de base de dados.xlsx')
sheet = workbook.active

# Diretório onde você deseja salvar as cartas
diretorio_saida = 'C:/Users/leonardo.farias/Desktop/Projetos/Gerador de Cartas/Cartas Geradas'

# Alterar pelas linhas da planilha (começando da segunda linha, pois a primeira linha contém cabeçalhos)
for row in sheet.iter_rows(min_row=2, values_only=True):

    # Carregar o modelo de carta
    template = Document('C:/Users/leonardo.farias/Desktop/Projetos/Gerador de Cartas/Carta de Reequilíbrio ContratualV2.docx')

    # Dicionário de substituições
    substituicoes = {
        'ID': str(row[0]),  # Substituir "ID" pelo valor da primeira coluna da planilha
        'EmbRot': f"R$ {locale.format_string('%0.2f', row[1], grouping=True)}" if row[1] != 0 else "Não Contratado",  # Substituir "Nome" pelo valor da segunda coluna da planilha
        'EmbEve': f"R$ {locale.format_string('%0.2f', row[2], grouping=True)}" if row[2] != 0 else "Não Contratado",  # Substituir "Valor1" pelo valor da terceira coluna da planilha
        'EmbEsp': f"R$ {locale.format_string('%0.2f', row[3], grouping=True)}" if row[3] != 0 else "Não Contratado",  # Substituir "Valor2" pelo valor da quarta coluna da planilha
        'FixoMensal': f"R$ {locale.format_string('%0.2f', row[4], grouping=True)}" if row[4] != 0 else "Não Contratado",  # Substituir zero por "sem alteração"
        'MilheCd': f"R$ {locale.format_string('%0.2f', row[5], grouping=True)}" if row[5] != 0 else "Não Contratado",
        'MilheMd': f"R$ {locale.format_string('%0.2f', row[6], grouping=True)}" if row[6] != 0 else "Não Contratado",
        'NomeCliente': str(row[7]),
        'DataHoje': data_extenso
        # Adicione mais substituições para outras palavras no modelo de carta, se necessário
    }

    # Realizar as substituições no modelo de carta
    for paragraph in template.paragraphs:
        for key, value in substituicoes.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    # Realizar as substituições nas tabelas do modelo de carta
    for table in template.tables:
        for row_table in table.rows:
            for cell in row_table.cells:
                for paragraph in cell.paragraphs:
                    for key, value in substituicoes.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(key, value)

    # Salvar a carta com o nome do cliente
    ID_cliente = row[0]
    Nome_cliente = row[7]
    nome_arquivo = f"Carta De Reajuste {Nome_cliente} - {ID_cliente}.docx"
    template.save(diretorio_saida + '/' + nome_arquivo)

    # Converter o documento DOCX para PDF
    convert(diretorio_saida + '/' + nome_arquivo)

print("Cartas criadas e convertidas para PDF com sucesso.")